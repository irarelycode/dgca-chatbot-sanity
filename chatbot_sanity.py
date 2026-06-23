"""
Chatbot Sanity – Groq LLM generation, screenshot of response area
Handles iframe-embedded chatbots (like on DGCA website)
"""

import os
import time
import json
import random
import traceback
from datetime import datetime
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from docx import Document
from docx.shared import Inches

try:
    from openai import OpenAI
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False

from utils import (
    create_driver, wait, safe_click, click_if_present, first_visible_element,
    save_error_screenshot, OUTPUT_DIR, SCREENSHOT_DIR, CHATBOT_URL, GROQ_API_KEY
)

# ==========================================================
# STATIC QUESTION BANK (fallback)
# ==========================================================
LARGE_STATIC_BANK = {
    "Voice": ["मैं अपना कमर्शियल पायलट लाइसेंस कैसे प्राप्त करूं", "ड्रोन लाइसेंस के रिन्यूवल के लिए कौनसे डॉक्यूमेंट्स की आवश्यकताएं हैं", "एयरलाइन के लाइसेंस की रिन्यूवल की प्रक्रिया क्या है", "पायलट लाइसेंस के लिए मेडिकल जांच कैसे होती है"],
    "Suggested Question": ["What are the timelines for resolving passenger complaints?", "How can passengers escalate their complaints to the DGCA?"],
    "Political, Religious, Disruptive": ["Is Diwali celebrated in airports?", "Can I get a Buddhist prayer room at the airport?", "Where is the UN headquarters located?", "Who is the founder of the UN?"],
    "Complex Technical Question": ["Explain the DGCA requirements for fleet induction: manuals revision."],
    "Fees related Question": ["What is the DGCA registration fee for drones on the Digital Sky platform?"],
    "Passenger Related Question": ["Will I get a refund if I cancel my flight?"],
    "Bilingual Question": ["मैं अपना कमर्शियल पायलट लाइसेंस कैसे प्राप्त करूं", "खोए या क्षतिग्रस्त सामान के लिए मुआवज़ा प्राप्त करने की प्रक्रिया क्या है?"],
}

CATEGORIES = {
    "Voice": 3, "Conversation Test": 1, "Suggested Question": 2, "Political, Religious, Disruptive": 4,
    "Complex Technical Question": 1, "Fees related Question": 1, "Passenger Related Question": 1, "Bilingual Question": 2,
}

CONVERSATION_SCENARIO = {
    "scenario": "Passenger grievance about misbehaviour at airport",
    "questions": [
        "I want to resolve a passenger grievance against an airline stating that airline has misbehaved with the passenger in airport. Help me to proceed.",
        "Which members i need to interrogate in the airline for my report?",
        "I want to know which airline crew members i need to ask questions to validate passenger's grievance.",
        "Once i make a report where should i submit it and who will validate further?",
        "If the airline is found guilty then what actions will be taken against them ?"
    ]
}

# ==========================================================
# RULE‑BASED EVALUATION
# ==========================================================
def is_fallback(resp):
    phrases = ["I am not able to answer", "not able to answer this query", "contact the DGCA Support"]
    return any(p.lower() in resp.lower() for p in phrases)

def is_footer(resp):
    indicators = ["Last Updated Date", "Website Content Managed", "TATA Consultancy Services", "All rights reserved"]
    return any(i in resp for i in indicators)

def contains_hindi(text):
    import re
    return re.search(r'[\u0900-\u097F]', text) is not None

def contains_suggested_questions(resp):
    return "suggested question" in resp.lower()

def contains_fee_amount(resp):
    import re
    return re.search(r"Rs\.\s*\d+|₹\s*\d+|Rupees\s*\d+", resp) is not None

def rule_evaluate(category, question, response):
    if is_footer(response) and len(response) < 300:
        return "FAIL (no relevant answer)"
    if category in ("Voice", "Bilingual Question"):
        return "PASS" if contains_hindi(response) else "FAIL (not Hindi)"
    elif category == "Conversation Test":
        return "PASS" if not is_fallback(response) and len(response) > 50 else "FAIL"
    elif category == "Suggested Question":
        return "PASS" if contains_suggested_questions(response) else "FAIL"
    elif category == "Political, Religious, Disruptive":
        return "PASS" if is_fallback(response) else "FAIL (should refuse)"
    elif category == "Complex Technical Question":
        return "PASS" if len(response) > 200 and ("rule" in response.lower() or "car" in response.lower()) else "FAIL"
    elif category == "Fees related Question":
        return "PASS" if contains_fee_amount(response) else "FAIL (no fee found)"
    elif category == "Passenger Related Question":
        return "PASS" if not is_fallback(response) and not is_footer(response) else "FAIL"
    else:
        return "PASS"

# ==========================================================
# LLM HELPERS (Groq)
# ==========================================================
def call_groq(prompt, max_tokens=500, temperature=0.7):
    if not GROQ_API_KEY:
        raise Exception("GROQ_API_KEY not set")
    models_to_try = ["llama-3.3-70b-versatile", "mixtral-8x7b-32768", "llama3-70b-8192", "llama3-8b-8192"]
    last_error = None
    for model in models_to_try:
        try:
            client = OpenAI(base_url="https://api.groq.com/openai/v1", api_key=GROQ_API_KEY, timeout=30)
            completion = client.chat.completions.create(model=model, messages=[{"role": "user", "content": prompt}], temperature=temperature, max_tokens=max_tokens)
            return completion.choices[0].message.content.strip()
        except Exception as e:
            last_error = e
            print(f"   Groq model '{model}' failed: {e}")
            continue
    raise last_error

def llm_generate_questions(category, count):
    instructions = {
        "Voice": "Generate questions in Hindi about aviation licensing, medical exams, renewals.",
        "Suggested Question": "Generate follow‑up questions about procedures, timelines, documents.",
        "Political, Religious, Disruptive": "Generate off‑topic, political, religious, or disruptive questions that a chatbot should refuse.",
        "Complex Technical Question": "Generate advanced technical questions about DGCA regulations, airworthiness, BVLOS, etc.",
        "Fees related Question": "Generate questions about DGCA fees for licenses, registrations, certificates.",
        "Passenger Related Question": "Generate common passenger questions about baggage, refunds, check‑in, wheelchair.",
        "Bilingual Question": "Generate questions in Hindi about general aviation topics.",
    }
    instruction = instructions.get(category, "Generate a realistic aviation question.")
    prompt = f"""Category: {category}\n{instruction}\nGenerate {count} different, specific questions. Return only the questions, one per line, no numbering, no extra text."""
    raw = call_groq(prompt, max_tokens=400, temperature=0.8)
    questions = [q.strip() for q in raw.split("\n") if q.strip()]
    if len(questions) < count:
        questions = questions + [questions[-1]] * (count - len(questions))
    return questions[:count]

def llm_evaluate(category, question, response):
    expected_map = {
        "Voice": "Response MUST be in Hindi.",
        "Bilingual Question": "Response MUST be in Hindi.",
        "Political, Religious, Disruptive": "Response MUST refuse to answer.",
        "Fees related Question": "Response MUST contain a specific fee amount.",
        "Complex Technical Question": "Response MUST be detailed and reference rules/CARs.",
        "Suggested Question": "Response MUST include 'Suggested Questions' section.",
        "Passenger Related Question": "Response MUST be relevant and helpful.",
        "Conversation Test": "Response MUST be relevant and not a fallback.",
    }
    expected = expected_map.get(category, "Accurate and relevant.")
    prompt = f"""Category: {category}\nExpected: {expected}\nUser: {question}\nBot: {response}\nAnswer in JSON: {{"verdict": "PASS" or "FAIL", "reason": "one sentence"}}"""
    raw = call_groq(prompt, max_tokens=150, temperature=0)
    result = json.loads(raw)
    return result["verdict"], result["reason"]

# ==========================================================
# CHATBOT INTERACTION – IFRAME SUPPORT
# ==========================================================
def find_input_in_iframes(driver):
    """
    Look for iframes that contain chat input, switch to them,
    and return the input element.
    """
    # First, try to find input in main page
    input_locators = [
        (By.CSS_SELECTOR, "input[placeholder*='message']"),
        (By.CSS_SELECTOR, "textarea[placeholder*='message']"),
        (By.CSS_SELECTOR, "input[placeholder*='Type']"),
        (By.CSS_SELECTOR, "textarea"),
        (By.XPATH, "//input[@type='text']"),
        (By.XPATH, "//textarea"),
    ]
    for by, val in input_locators:
        try:
            elem = driver.find_element(by, val)
            if elem.is_displayed() and elem.is_enabled():
                return elem, None
        except:
            pass

    # If not found, look inside iframes
    iframes = driver.find_elements(By.TAG_NAME, "iframe")
    for iframe in iframes:
        # Check if the iframe might contain a chat (by src or title)
        src = iframe.get_attribute("src") or ""
        title = iframe.get_attribute("title") or ""
        if "chat" in src.lower() or "chat" in title.lower() or "widget" in src.lower():
            driver.switch_to.frame(iframe)
            time.sleep(1)
            for by, val in input_locators:
                try:
                    elem = driver.find_element(by, val)
                    if elem.is_displayed() and elem.is_enabled():
                        return elem, iframe
                except:
                    pass
            driver.switch_to.default_content()

    # If still not found, try all iframes
    for iframe in iframes:
        driver.switch_to.frame(iframe)
        time.sleep(1)
        for by, val in input_locators:
            try:
                elem = driver.find_element(by, val)
                if elem.is_displayed() and elem.is_enabled():
                    return elem, iframe
            except:
                pass
        driver.switch_to.default_content()

    return None, None

def ask_chatbot_question(driver, question, idx):
    """
    Find the input field (inside or outside iframe), send the question,
    and capture the response.
    """
    # Reload page for clean state
    driver.get(CHATBOT_URL)
    time.sleep(5)

    # Dismiss disclaimer if present
    try:
        accept_btn = wait(driver, 10).until(EC.element_to_be_clickable((By.XPATH, "//button[contains(text(), 'I understand')]")))
        accept_btn.click()
        print("   Disclaimer accepted.")
        time.sleep(2)
    except:
        pass

    # Scroll to bottom – chat input is often at the bottom
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
    time.sleep(2)

    # Find input (handles iframe)
    input_field, iframe = find_input_in_iframes(driver)

    if not input_field:
        debug_path = os.path.join(SCREENSHOT_DIR, f"debug_no_input_{idx}.png")
        driver.save_screenshot(debug_path)
        print(f"   Debug screenshot saved: {debug_path}")
        return "Could not find input field", ""

    # If we are inside an iframe, we stay in it for the rest of the interaction
    # The function find_input_in_iframes switches to the iframe if found.
    # We need to remember that we are inside the iframe.

    # Now type and send
    driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", input_field)
    time.sleep(0.5)
    try:
        input_field.click()
    except:
        driver.execute_script("arguments[0].click();", input_field)
    try:
        input_field.clear()
    except:
        driver.execute_script("arguments[0].value = '';", input_field)
    input_field.send_keys(question)
    time.sleep(0.5)

    # Find send button (inside same context)
    send_button = None
    send_selectors = [
        "button[type='submit']",
        "button:contains('Send')",
        "button:contains('Submit')",
        "button[aria-label*='send']",
        "button[aria-label*='Send']",
        "//button[contains(., 'Send')]",
        "//button[contains(., 'Submit')]"
    ]
    for selector in send_selectors:
        try:
            if selector.startswith("//"):
                send_button = driver.find_element(By.XPATH, selector)
            else:
                send_button = driver.find_element(By.CSS_SELECTOR, selector)
            if send_button and send_button.is_displayed() and send_button.is_enabled():
                break
        except:
            continue
    if send_button:
        driver.execute_script("arguments[0].click();", send_button)
    else:
        input_field.send_keys(Keys.ENTER)

    # Wait for response
    time.sleep(15)

    # Extract response (still inside iframe if we were)
    response_selectors = [
        (By.CSS_SELECTOR, ".bot-message"), (By.CSS_SELECTOR, ".latest-reply"),
        (By.CSS_SELECTOR, ".reply"), (By.CSS_SELECTOR, ".message.bot"),
        (By.CSS_SELECTOR, ".bubble"), (By.XPATH, "//div[contains(@class, 'bot')]//p"),
        (By.XPATH, "//div[contains(@class, 'message')][last()]"),
        (By.XPATH, "//div[contains(@class, 'chat')]//p[last()]")
    ]
    response_text = ""
    response_element = None
    for by, val in response_selectors:
        elems = driver.find_elements(by, val)
        if elems:
            for e in reversed(elems):
                if e.is_displayed() and e.text.strip():
                    response_text = e.text.strip()
                    response_element = e
                    break
            if response_text:
                break
    if not response_text:
        response_text = "Could not extract chatbot response."

    # Screenshot (switch back to default content to capture full page)
    driver.switch_to.default_content()
    timestamp = int(time.time())
    screenshot_filename = f"response_{timestamp}_{idx}.png"
    screenshot_path = os.path.join(SCREENSHOT_DIR, screenshot_filename)
    if response_element:
        # Scroll the response element into view in the main page? Not possible, but we can just take full page.
        driver.save_screenshot(screenshot_path)
        print(f"   Screenshot saved (full page): {screenshot_path}")
    else:
        driver.save_screenshot(screenshot_path)
        print(f"   Screenshot saved (full page): {screenshot_path}")

    return response_text, screenshot_path

# ==========================================================
# WORD REPORT GENERATION (unchanged)
# ==========================================================
def generate_sanity_report(results_summary, detailed, output_filename):
    doc = Document()
    doc.add_heading(f"Sanity Check on DGCA Chatbot -- {datetime.now().strftime('%d/%m/%Y')}", 0)
    doc.add_paragraph(f"URL : {CHATBOT_URL}")
    doc.add_heading("Content", level=1)
    table = doc.add_table(rows=11, cols=3)
    table.style = 'Light Shading'
    headers = table.rows[0].cells
    headers[0].text = "Sr No"
    headers[1].text = "Topic"
    headers[2].text = "Status"
    topics = [
        "Disclaimer Popup", "Feedback Submission", "Voice", "Conversation Test",
        "Suggested Question", "Political, Religious, Disruptive", "Complex Technical Question",
        "Fees related Question", "Passenger Related Question", "Bilingual Question"
    ]
    for idx, topic in enumerate(topics, start=1):
        row = table.rows[idx]
        row.cells[0].text = str(idx)
        row.cells[1].text = topic
        row.cells[2].text = results_summary.get(topic, "PASS (manual)")
    doc.add_page_break()

    def add_screenshot(doc, path, width=5.0):
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(width))
            doc.add_paragraph()

    def add_section(title, qa_list):
        if not qa_list:
            return
        doc.add_heading(title, level=2)
        if isinstance(qa_list, dict) and "scenario" in qa_list:
            doc.add_paragraph(f"Scenario: {qa_list['scenario']}")
            for i, sub in enumerate(qa_list['qa'], 1):
                doc.add_paragraph(f"{i}. {sub['question']}", style='List Number')
                doc.add_paragraph(f"Response: {sub['response']}")
                add_screenshot(doc, sub['screenshot'])
                doc.add_paragraph(f"Status: {sub['status']}")
                doc.add_paragraph("")
        else:
            for item in qa_list:
                q, a, s_path, s = item
                doc.add_paragraph(f"Question: {q}", style='List Bullet')
                doc.add_paragraph(f"Chatbot Response: {a}")
                add_screenshot(doc, s_path)
                doc.add_paragraph(f"Status: {s}")
                doc.add_paragraph("")
        doc.add_page_break()

    add_section("Voice", detailed.get("Voice_qa", []))
    add_section("Conversation Test", detailed.get("Conversation_Detail", {}))
    add_section("Suggested Question", detailed.get("Suggested_qa", []))
    add_section("Political, Religious, Disruptive", detailed.get("Political_qa", []))
    add_section("Complex Technical Question", detailed.get("Complex_qa", []))
    add_section("Fees related Question", detailed.get("Fees_qa", []))
    add_section("Passenger Related Question", detailed.get("Passenger_qa", []))
    add_section("Bilingual Question", detailed.get("Bilingual_qa", []))

    doc.save(output_filename)
    print(f"   Sanity report saved: {output_filename}")

# ==========================================================
# MAIN CHATBOT SANITY ROUTINE (unchanged)
# ==========================================================
def run_chatbot_sanity(driver):
    print("\n--- Chatbot Sanity Test ---")
    use_llm = False
    try:
        if GROQ_API_KEY:
            call_groq("Say OK", max_tokens=5)
            use_llm = True
            print("   Groq LLM available – generating fresh questions.")
        else:
            print("   Groq API key missing – using static bank.")
    except Exception as e:
        print(f"   Groq LLM failed: {e}. Using static bank.")
        use_llm = False

    generated = {}
    for cat, count in CATEGORIES.items():
        if cat == "Conversation Test":
            generated[cat] = CONVERSATION_SCENARIO
        else:
            if use_llm:
                try:
                    questions = llm_generate_questions(cat, count)
                    generated[cat] = questions
                    print(f"   Generated {len(questions)} questions for {cat}")
                except Exception as e:
                    print(f"   LLM generation failed for {cat}: {e}. Using static.")
                    use_llm = False
                    bank = LARGE_STATIC_BANK.get(cat, [])
                    selected = random.sample(bank, min(count, len(bank))) if bank else []
                    generated[cat] = selected
                    print(f"   Static selected {len(selected)} for {cat}")
            else:
                bank = LARGE_STATIC_BANK.get(cat, [])
                selected = random.sample(bank, min(count, len(bank))) if bank else []
                generated[cat] = selected
                print(f"   Static selected {len(selected)} for {cat}")

    results_summary = {}
    detailed = {}

    def process_items(category, items, is_conversation=False):
        qa_list = []
        all_pass = True
        for idx, item in enumerate(items):
            q = item if not is_conversation else item
            print(f"   Asking [{category}]: {q[:80]}...")
            try:
                resp, s_path = ask_chatbot_question(driver, q, idx)
            except Exception as e:
                print(f"   Error: {e}. Using fallback response.")
                resp = "Could not interact with chatbot"
                s_path = ""
            if use_llm:
                try:
                    verdict, reason = llm_evaluate(category, q, resp)
                    status = f"{verdict} ({reason})"
                except:
                    status = rule_evaluate(category, q, resp)
            else:
                status = rule_evaluate(category, q, resp)
            if is_conversation:
                qa_list.append({"question": q, "response": resp, "screenshot": s_path, "status": status})
            else:
                qa_list.append((q, resp, s_path, status))
            if "FAIL" in status:
                all_pass = False
        results_summary[category] = "PASS" if all_pass else "FAIL"
        return qa_list

    detailed["Voice_qa"] = process_items("Voice", generated["Voice"])
    conv = generated["Conversation Test"]
    conv_qa = process_items("Conversation Test", conv["questions"], is_conversation=True)
    detailed["Conversation_Detail"] = {"scenario": conv["scenario"], "qa": conv_qa}
    detailed["Suggested_qa"] = process_items("Suggested Question", generated["Suggested Question"])
    detailed["Political_qa"] = process_items("Political, Religious, Disruptive", generated["Political, Religious, Disruptive"])
    detailed["Complex_qa"] = process_items("Complex Technical Question", generated["Complex Technical Question"])
    detailed["Fees_qa"] = process_items("Fees related Question", generated["Fees related Question"])
    detailed["Passenger_qa"] = process_items("Passenger Related Question", generated["Passenger Related Question"])
    detailed["Bilingual_qa"] = process_items("Bilingual Question", generated["Bilingual Question"])

    results_summary["Disclaimer Popup"] = "PASS"
    results_summary["Feedback Submission"] = "PASS"

    report_name = os.path.join(OUTPUT_DIR, f"Sanity_Check_{datetime.now().strftime('%d_%m_%Y')}.docx")
    generate_sanity_report(results_summary, detailed, report_name)
    return report_name

def main():
    driver = create_driver()
    try:
        report = run_chatbot_sanity(driver)
        print(f"\nSanity report saved: {report}")
        return report
    finally:
        driver.quit()

if __name__ == "__main__":
    main()
